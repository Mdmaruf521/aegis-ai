"""
ATS Surgeon Pro v4 - Backend (FREE VERSION)
Uses Ollama (local LLM) -- no API key, no cost, runs on your Mac.

SETUP (one time):
  brew install ollama
  ollama pull llama3
  ollama serve          <- keep this running in a separate terminal

Then start this server:
  python3 -m uvicorn main:app --reload
"""

import re
import io
import json
import requests
import fitz  # PyMuPDF
import spacy
from typing import List, Dict
from fastapi import FastAPI, UploadFile, Form, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sentence_transformers import SentenceTransformer, util
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.enums import TA_CENTER

# Config
OLLAMA_URL   = "http://localhost:11434/api/generate"
OLLAMA_MODEL = "llama3"   # change to "mistral" or "phi3" if preferred

# App init
app = FastAPI(title="ATS Surgeon Pro v4 (Free/Local)")
embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
nlp = spacy.load("en_core_web_sm")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], allow_methods=["*"], allow_headers=["*"],
)

# Pydantic models
class BulletAnalysis(BaseModel):
    original: str
    issue: str
    fix: str

class SectionScore(BaseModel):
    score: float
    status: str
    feedback: str
    key_gaps: List[str]
    bullet_fixes: List[BulletAnalysis]

class KeywordDetail(BaseModel):
    keyword: str
    present: bool
    importance: str
    context: str

class AnalysisResponse(BaseModel):
    overall_score: float
    ats_compatibility: float
    impact_score: float
    sections: Dict[str, SectionScore]
    keyword_analysis: List[KeywordDetail]
    missing_keywords: List[str]
    present_keywords: List[str]
    critical_issues: List[str]
    quick_wins: List[str]
    optimization_strategy: str
    rewrite_available: bool

# Ollama helper
def call_ollama(prompt: str) -> str:
    try:
        resp = requests.post(
            "http://localhost:11434/api/chat",
            json={
                "model": OLLAMA_MODEL,
                "messages": [
                    {"role": "user", "content": prompt}
                ],
                "stream": False,
                "format": "json"   # 🔥 IMPORTANT FIX (forces structured JSON output)
            },
            timeout=180,
        )

        resp.raise_for_status()
        data = resp.json()

        # Ollama chat response safety extraction
        if "message" in data and "content" in data["message"]:
            return data["message"]["content"]

        raise ValueError(f"Unexpected Ollama response: {data}")

    except requests.exceptions.ConnectionError:
        raise RuntimeError("Ollama not running: run `ollama serve`")
    
def parse_json_from_llm(raw: str) -> dict:
    if not raw:
        raise ValueError("Empty LLM response")

    # Remove markdown noise
    raw = re.sub(r"```(?:json)?", "", raw).strip().strip("`").strip()

    # Try direct parse first (best case)
    try:
        return json.loads(raw)
    except:
        pass

    # Fallback: extract JSON block
    start = raw.find("{")
    end = raw.rfind("}") + 1

    if start == -1 or end == 0:
        raise ValueError(f"No JSON found in LLM output: {raw[:200]}")

    try:
        return json.loads(raw[start:end])
    except json.JSONDecodeError as e:
        raise ValueError(f"Malformed JSON from LLM: {str(e)}\nRAW: {raw[:300]}")

# Text helpers
def extract_text_from_pdf(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    return "\n".join(page.get_text() for page in doc)

def segment_resume(text: str) -> Dict[str, str]:
    headers = {
        "experience": r"(experience|work history|employment|career|professional background)",
        "education":  r"(education|academic|qualification|degree|studies|university|college)",
        "skills":     r"(skills|technical skills|core competencies|technologies|stack|tools|expertise)",
        "projects":   r"(projects|portfolio|personal work|key projects|notable projects)",
        "awards":     r"(awards|honors|achievements|recognition|certifications|certificates|accomplishments)",
        "summary":    r"(summary|profile|objective|about|overview|professional summary)",
    }
    segments = {k: "" for k in headers}
    current = None
    for line in text.split("\n"):
        clean = line.strip().lower()
        matched = False
        for key, pattern in headers.items():
            if re.search(r"^" + pattern, clean):
                current = key; matched = True; break
        if not matched and current:
            segments[current] += line + "\n"
    return segments

def extract_keywords_advanced(text: str, limit: int = 20) -> List[str]:
    tech_pattern = re.compile(
        r"\b(python|java|javascript|typescript|react|node|sql|aws|docker|kubernetes|"
        r"machine learning|deep learning|nlp|api|rest|graphql|git|agile|scrum|"
        r"tensorflow|pytorch|spark|hadoop|excel|tableau|power bi|c\+\+|c#|\.net|"
        r"azure|gcp|redis|mongodb|postgresql|mysql|kafka|airflow|fastapi|django|flask|"
        r"ci/cd|devops|microservices|llm|ai|ml|data science|analytics|leadership|"
        r"communication|teamwork|problem.solving|project management)\b",
        re.IGNORECASE,
    )
    tech_kws = list(set(tech_pattern.findall(text.lower())))
    doc = nlp(text.lower())
    noun_kws = [t.lemma_ for t in doc if t.pos_ in ("NOUN","PROPN") and not t.is_stop and len(t.text) > 2]
    return list(set(tech_kws + noun_kws))[:limit]

def calculate_ats_compatibility(cv_text: str) -> float:
    score = 100.0
    if not re.search(r"\b\d{4}\b", cv_text):                                         score -= 10
    if not re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", cv_text):  score -= 15
    if not re.search(r"linkedin\.com|github\.com", cv_text.lower()):                 score -= 5
    if len(re.findall(r"^[\u2022\-\*\-]", cv_text, re.MULTILINE)) < 5:                 score -= 10
    if len(re.findall(r"%|\$|increased|decreased|improved|reduced|delivered", cv_text.lower())) < 2: score -= 10
    if len(cv_text) < 500: score -= 20
    if re.search(r"\t{3,}", cv_text): score -= 15
    return max(0.0, round(score, 1))

def calculate_impact_score(cv_text: str) -> float:
    score = 50.0
    metrics = len(re.findall(r"\d+[\%\+]?|\$[\d,]+|[\d,]+\s*(users|clients|revenue|records|percent)", cv_text.lower()))
    score += min(metrics * 5, 25)
    verbs = ["led","built","designed","architected","scaled","launched","delivered",
             "managed","increased","reduced","improved","developed","created","implemented","optimized"]
    score += min(sum(1 for v in verbs if re.search(r"\b"+v+r"\b", cv_text.lower())) * 2, 20)
    if re.search(r"(award|recognition|honor|promoted|distinction)", cv_text.lower()): score += 5
    return min(100.0, round(score, 1))

FALLBACK_ANALYSIS = {
    "critical_issues": ["LLM analysis failed (format/parsing issue). Ollama is running but response was not strict JSON."],
    "quick_wins": ["Run: ollama serve in a separate terminal, then retry"],
    "optimization_strategy": "Scores above are based on semantic similarity only. Start Ollama for AI-powered analysis.",
    "section_feedback": {s: {"key_gaps":[],"bullet_fixes":[]} for s in ["experience","skills","education","projects","awards","summary"]},
    "keyword_analysis": []
}

def deep_analyze_with_llm(cv_text: str, jd: str) -> Dict:
    prompt = (
        "You are an expert ATS resume analyst. Analyze the resume against the job description.\n\n"
        f"JOB DESCRIPTION:\n{jd[:2500]}\n\nRESUME:\n{cv_text[:3500]}\n\n"
        "Respond with ONLY a raw JSON object. No markdown, no explanation, no backticks. Use this exact structure:\n"
        '{"critical_issues":["issue1","issue2"],"quick_wins":["win1","win2"],"optimization_strategy":"3-4 sentence overview",'
        '"section_feedback":{"experience":{"key_gaps":["gap1"],"bullet_fixes":[{"original":"text","issue":"issue","fix":"fix"}]},'
        '"skills":{"key_gaps":["gap1"],"bullet_fixes":[]},"education":{"key_gaps":[],"bullet_fixes":[]},'
        '"projects":{"key_gaps":["gap1"],"bullet_fixes":[{"original":"text","issue":"issue","fix":"fix"}]},'
        '"awards":{"key_gaps":[],"bullet_fixes":[]},"summary":{"key_gaps":["gap1"],"bullet_fixes":[{"original":"text","issue":"issue","fix":"fix"}]}},'
        '"keyword_analysis":[{"keyword":"python","present":true,"importance":"critical","context":"why it matters"},'
        '{"keyword":"kubernetes","present":false,"importance":"high","context":"why it matters"}]}'
    )
    try:
        raw = call_ollama(prompt)
        return parse_json_from_llm(raw)
    except Exception as e:
        print(f"[WARN] LLM analysis failed: {e}")
        return FALLBACK_ANALYSIS

def rewrite_resume_with_llm(cv_text: str, jd: str) -> Dict:
    prompt = (
        "You are a professional resume writer. Rewrite this resume to match the job description. Keep all facts accurate.\n\n"
        f"JOB DESCRIPTION:\n{jd[:2500]}\n\nORIGINAL RESUME:\n{cv_text[:3500]}\n\n"
        "Respond with ONLY a raw JSON object. No markdown, no backticks:\n"
        '{"name":"Full Name","contact":"email | phone | linkedin | location",'
        '"summary":"2-3 sentence professional summary",'
        '"experience":[{"title":"Job Title","company":"Company","dates":"Start-End","bullets":["bullet1","bullet2"]}],'
        '"skills":{"technical":["skill1","skill2"],"soft":["skill1"]},'
        '"education":[{"degree":"Degree","institution":"School","year":"Year","details":"details"}],'
        '"projects":[{"name":"Project","description":"description","tech":["tech1"]}],'
        '"awards":["award1"]}'
    )
    raw = call_ollama(prompt)
    return parse_json_from_llm(raw)

def build_pdf(rd: Dict) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            rightMargin=0.75*inch, leftMargin=0.75*inch,
                            topMargin=0.75*inch, bottomMargin=0.75*inch)
    def s(**kw): return ParagraphStyle("x", **kw)
    ns  = s(fontSize=22, fontName="Helvetica-Bold", textColor=colors.HexColor("#1a1a2e"), alignment=TA_CENTER, spaceAfter=4)
    cs  = s(fontSize=9,  fontName="Helvetica",      textColor=colors.HexColor("#555555"), alignment=TA_CENTER, spaceAfter=12)
    ss  = s(fontSize=11, fontName="Helvetica-Bold", textColor=colors.HexColor("#1a56db"), spaceBefore=12, spaceAfter=4)
    bs  = s(fontSize=9.5,fontName="Helvetica",      textColor=colors.HexColor("#333333"), spaceAfter=3, leading=14)
    bls = s(fontSize=9.5,fontName="Helvetica",      textColor=colors.HexColor("#333333"), leftIndent=12, spaceAfter=2, leading=14)
    js  = s(fontSize=10, fontName="Helvetica-Bold", textColor=colors.HexColor("#222222"), spaceAfter=1)
    cos = s(fontSize=9.5,fontName="Helvetica-Oblique",textColor=colors.HexColor("#555555"), spaceAfter=2)
    story = [Paragraph(rd.get("name","Candidate"), ns), Paragraph(rd.get("contact",""), cs),
             HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#1a56db"))]
    def sec(t):
        story.extend([Paragraph(t, ss), HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#dddddd")), Spacer(1,4)])
    if rd.get("summary"): sec("PROFESSIONAL SUMMARY"); story.append(Paragraph(rd["summary"], bs))
    if rd.get("experience"):
        sec("PROFESSIONAL EXPERIENCE")
        for exp in rd["experience"]:
            story += [Spacer(1,6), Paragraph(exp.get("title",""), js),
                      Paragraph(f'{exp.get("company","")}  |  {exp.get("dates","")}', cos)]
            for b in exp.get("bullets",[]): story.append(Paragraph(f"\u2022 {b}", bls))
    if rd.get("skills"):
        sec("SKILLS"); sk = rd["skills"]
        if isinstance(sk,dict):
            if sk.get("technical"): story.append(Paragraph(f'<b>Technical:</b> {", ".join(sk["technical"])}', bs))
            if sk.get("soft"):      story.append(Paragraph(f'<b>Soft Skills:</b> {", ".join(sk["soft"])}', bs))
        else: story.append(Paragraph(", ".join(sk), bs))
    if rd.get("projects"):
        sec("PROJECTS")
        for p in rd["projects"]:
            tech = ", ".join(p.get("tech",[]))
            story += [Spacer(1,6), Paragraph(f'<b>{p.get("name","")}</b>  <font size=8 color="#666">({tech})</font>', bs)]
            desc = p.get("description","")
            for line in (desc if isinstance(desc,list) else [desc]): story.append(Paragraph(f"\u2022 {line}", bls))
    if rd.get("education"):
        sec("EDUCATION")
        for edu in rd["education"]:
            story += [Spacer(1,4), Paragraph(f'<b>{edu.get("degree","")}</b>  ---  {edu.get("institution","")}, {edu.get("year","")}', bs)]
            if edu.get("details"): story.append(Paragraph(edu["details"], bls))
    if rd.get("awards"):
        sec("AWARDS & CERTIFICATIONS")
        for a in rd["awards"]: story.append(Paragraph(f"\u2022 {a}", bls))
    doc.build(story); return buf.getvalue()

def build_docx(rd: Dict) -> bytes:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(0.75)
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10)
    def add_heading(title):
        p = doc.add_paragraph(); r = p.add_run(title); r.bold = True; r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(26,86,219); p.paragraph_format.space_before = Pt(12)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(rd.get("name","Candidate")); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(26,26,46)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(rd.get("contact","")); r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(85,85,85)
    if rd.get("summary"): add_heading("PROFESSIONAL SUMMARY"); doc.add_paragraph(rd["summary"])
    if rd.get("experience"):
        add_heading("PROFESSIONAL EXPERIENCE")
        for exp in rd["experience"]:
            p = doc.add_paragraph(); r = p.add_run(exp.get("title","")); r.bold = True; r.font.size = Pt(10.5)
            p2 = doc.add_paragraph(); r2 = p2.add_run(f'{exp.get("company","")}  |  {exp.get("dates","")}'  )
            r2.italic = True; r2.font.color.rgb = RGBColor(85,85,85); r2.font.size = Pt(9.5)
            for b in exp.get("bullets",[]): doc.add_paragraph(b, style="List Bullet")
    if rd.get("skills"):
        add_heading("SKILLS"); sk = rd["skills"]
        if isinstance(sk,dict):
            if sk.get("technical"):
                p = doc.add_paragraph(); p.add_run("Technical: ").bold = True; p.add_run(", ".join(sk["technical"]))
            if sk.get("soft"):
                p = doc.add_paragraph(); p.add_run("Soft Skills: ").bold = True; p.add_run(", ".join(sk["soft"]))
        else: doc.add_paragraph(", ".join(sk))
    if rd.get("projects"):
        add_heading("PROJECTS")
        for proj in rd["projects"]:
            p = doc.add_paragraph(); r = p.add_run(proj.get("name","")); r.bold = True
            tech = ", ".join(proj.get("tech",[]));
            if tech: p.add_run(f"  ({tech})")
            desc = proj.get("description","")
            for line in (desc if isinstance(desc,list) else [desc]): doc.add_paragraph(line, style="List Bullet")
    if rd.get("education"):
        add_heading("EDUCATION")
        for edu in rd["education"]:
            p = doc.add_paragraph(); r = p.add_run(f'{edu.get("degree","")}  ---  '); r.bold = True
            p.add_run(f'{edu.get("institution","")}, {edu.get("year","")}'  )
            if edu.get("details"): doc.add_paragraph(edu["details"])
    if rd.get("awards"):
        add_heading("AWARDS & CERTIFICATIONS")
        for a in rd["awards"]: doc.add_paragraph(a, style="List Bullet")
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

@app.post("/v4/analyze", response_model=AnalysisResponse)
def analyze_v4(cv: UploadFile = File(...), jd: str = Form(...)):
    try:
        raw_text = extract_text_from_pdf(cv.file.read())
        if not raw_text.strip(): raise HTTPException(400, "PDF is empty or unreadable.")
        segments = segment_resume(raw_text)
        jd_kws   = extract_keywords_advanced(jd, 25)
        cv_kws   = extract_keywords_advanced(raw_text, 30)
        cv_lower = [k.lower() for k in cv_kws]
        missing  = [k for k in jd_kws if k.lower() not in cv_lower]
        present  = [k for k in jd_kws if k.lower() in cv_lower]
        weights  = {"experience":0.40,"skills":0.30,"projects":0.12,"education":0.08,"awards":0.05,"summary":0.05}
        section_results = {}; total_score = 0.0
        for sec, weight in weights.items():
            sec_text = segments.get(sec,"").strip(); score = 0.0
            if sec_text:
                emb   = embedding_model.encode([sec_text, jd])
                score = round(float(util.cos_sim(emb[0], emb[1]).item()) * 100, 2)
            status   = "STRONG" if score>75 else "PASS" if score>55 else "REVISE" if score>30 else "FAIL"
            feedback = ("Excellent alignment." if score>75 else "Good alignment, minor gaps." if score>55
                        else "Significant improvements needed." if score>30 else "Critical gap -- needs rework.")
            section_results[sec] = {"score":score,"status":status,"feedback":feedback,"key_gaps":[],"bullet_fixes":[]}
            total_score += score * weight
        llm_data = deep_analyze_with_llm(raw_text, jd)
        for sec in section_results:
            sf = llm_data.get("section_feedback",{}).get(sec,{})
            section_results[sec]["key_gaps"]     = sf.get("key_gaps",[])
            section_results[sec]["bullet_fixes"] = sf.get("bullet_fixes",[])
        keyword_analysis = [
            KeywordDetail(keyword=k.get("keyword",""), present=k.get("present",False),
                          importance=k.get("importance","medium"), context=k.get("context",""))
            for k in llm_data.get("keyword_analysis",[])
        ]
        return AnalysisResponse(
            overall_score=round(total_score,1), ats_compatibility=calculate_ats_compatibility(raw_text),
            impact_score=calculate_impact_score(raw_text),
            sections={k: SectionScore(**v) for k,v in section_results.items()},
            keyword_analysis=keyword_analysis,
            missing_keywords=missing or ["No critical keywords missing"],
            present_keywords=present,
            critical_issues=llm_data.get("critical_issues",[]),
            quick_wins=llm_data.get("quick_wins",[]),
            optimization_strategy=llm_data.get("optimization_strategy",""),
            rewrite_available=True,
        )
    except HTTPException: raise
    except Exception as e: print(f"CRITICAL ERROR: {e}"); raise HTTPException(500, str(e))

@app.post("/v4/rewrite")
def rewrite_resume(cv: UploadFile = File(...), jd: str = Form(...), format: str = Form(default="pdf")):
    try:
        raw_text = extract_text_from_pdf(cv.file.read())
        if not raw_text.strip(): raise HTTPException(400, "PDF is empty or unreadable.")
        rd = rewrite_resume_with_llm(raw_text, jd)
        if format == "docx":
            return StreamingResponse(io.BytesIO(build_docx(rd)),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": "attachment; filename=optimized_resume.docx"})
        return StreamingResponse(io.BytesIO(build_pdf(rd)),
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=optimized_resume.pdf"})
    except HTTPException: raise
    except Exception as e: print(f"REWRITE ERROR: {e}"); raise HTTPException(500, str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)